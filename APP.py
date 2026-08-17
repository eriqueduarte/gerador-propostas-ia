import streamlit as st
import google.generativeai as genai
from docx import Document
import io

# Configuração da página da web
st.set_page_config(page_title="Gerador de Propostas IA", page_icon="💼", layout="centered")

st.title("💼 Gerador Inteligente de Propostas Comerciais")
st.write("Insira seus dados para gerar propostas comerciais altamente persuasivas usando a API do Gemini.")

# Campo seguro para colocar a chave que você pegou no Google AI Studio
api_key_input = st.text_input("Insira sua API Key do Google AI Studio:", type="password", 
                             help="Cole aqui a chave que você copiou do painel do Google.")

# Formulário de entrada de dados da proposta
with st.form("dados_proposta"):
    nome_cliente = st.text_input("Nome do Cliente / Empresa:", placeholder="Ex: Clínica OdontoSorriso")
    servico = st.text_input("Serviço Ofertado:", placeholder="Ex: Gestão de Tráfego Pago e Redes Sociais")
    dores_cliente = st.text_area("Quais os problemas/dores que o cliente enfrenta hoje?", 
                                 placeholder="Ex: Dependem apenas de indicações e o faturamento travou.")
    preco_estimado = st.text_input("Valor do Investimento / Condições:", placeholder="Ex: R$ 2.500/mês")
    
    botao_gerar = st.form_submit_button("Gerar Proposta Comercial 🔥")

# Função auxiliar para converter o texto gerado em um arquivo Word (.docx)
def criar_arquivo_word(texto_proposta):
    doc = Document()
    doc.add_heading('Proposta Comercial', level=1)
    
    # Divide o texto por linhas e adiciona ao documento
    for linha in texto_proposta.split('\n'):
        if linha.strip().startswith('**') and linha.strip().endswith('**'):
            doc.add_heading(linha.replace('**', ''), level=2)
        elif linha.strip().startswith('#'):
            doc.add_heading(linha.replace('#', '').strip(), level=2)
        else:
            doc.add_paragraph(linha.replace('**', '').replace('*', ''))
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Lógica de processamento
if botao_gerar:
    if not api_key_input:
        st.error("⚠️ Por favor, insira sua API Key no campo acima para prosseguir.")
    elif not nome_cliente or not servico or not dores_cliente or not preco_estimado:
        st.error("⚠️ Por favor, preencha todos os campos do formulário.")
    else:
        with st.spinner("O Gemini Pro está estruturando sua proposta perfeita..."):
            try:
                # Configura a API com a chave inserida pelo usuário
                genai.configure(api_key=api_key_input)
                
                # Inicializa o modelo do Gemini
                model = genai.GenerativeModel('gemini-3.5-flash')
                
                prompt = f"""
                Você é um copywriter profissional e especialista em vendas B2B de alto impacto.
                Escreva uma proposta comercial altamente persuasiva com base nos seguintes dados:
                - Cliente: {nome_cliente}
                - Serviço Prestado: {servico}
                - Principais problemas/dores que o cliente enfrenta: {dores_cliente}
                - Valor do Investimento: {preco_estimado}
                
                A proposta deve seguir a seguinte estrutura estrita em Markdown:
                1. **Sumário Executivo**: Um gancho forte mostrando que entendemos a dor atual dele.
                2. **A Solução**: Como o nosso serviço resolve especificamente as dores mencionadas.
                3. **Escopo do Trabalho**: Entregáveis claros em tópicos.
                4. **Investimento e Condições**: O valor de {preco_estimado} detalhado e profissionalmente justificado como investimento (ROI).
                5. **Próximos Passos**: Chamada para ação clara para fechar o contrato.
                
                Use um tom profissional, corporativo, porém altamente focado em conversão e resultados. Não inclua saudações iniciais fora do Sumário.
                """
                
                response = model.generate_content(prompt)
                texto_gerado = response.text
                
                # Armazena o resultado no estado da sessão para não sumir ao clicar no botão de baixar
                st.session_state['proposta'] = texto_gerado
                st.session_state['nome_cliente'] = nome_cliente
                
            except Exception as e:
                st.error(f"Erro ao conectar com a API: {e}. Verifique se sua chave está correta.")

# Se a proposta já foi gerada, exibe na tela e mostra o botão de download
if 'proposta' in st.session_state:
    st.success("✨ Proposta gerada com sucesso!")
    st.markdown("---")
    st.markdown(st.session_state['proposta'])
    st.markdown("---")
    
    # Cria o arquivo na memória para o download
    arquivo_docx = criar_arquivo_word(st.session_state['proposta'])
    nome_arquivo = f"Proposta_{st.session_state['nome_cliente'].replace(' ', '_')}.docx"
    
    # Botão de download nativo do Streamlit
    st.download_button(
        label="📥 Baixar Proposta em Word (.docx)",
        data=arquivo_docx,
        file_name=nome_arquivo,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
