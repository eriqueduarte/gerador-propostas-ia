import streamlit as st
import google.generativeai as genai
from docx import Document
import io

# =========================================================================
# CONFIGURAÇÕES DE SEGURANÇA, NEGÓCIO E API (SUAS CONFIGURAÇÕES PRIVADAS)
# =========================================================================
SENHA_CORRETA = "PROPOSTA2026"  # Senha que você enviará aos compradores
LINK_MERCADO_PAGO = "https://mpago.la/1ud7mBi"  # Seu link do Mercado Pago

# Cole aqui a sua chave do Google AI Studio (Aquela que termina em ...6ftA)
MINHA_API_KEY_PRIVADA = st.secrets["GEMINI_API_KEY"]
# =========================================================================

st.set_page_config(page_title="Gerador de Propostas IA", page_icon="💼", layout="centered")

st.title("💼 Gerador Inteligente de Propostas Comerciais")

# Inicializa o estado da senha se não existir
if 'autenticado' not in st.session_state:
    st.session_state['autenticado'] = False

# Tela de Bloqueio (Se o usuário ainda não digitou a senha correta)
if not st.session_state['autenticado']:
    st.write("### 🔒 Área Restrita a Assinantes")
    st.write("Esta ferramenta automatizada cria propostas comerciais de alto impacto e exporta direto para o Word.")
    
    senha_digitada = st.text_input("Insira sua Senha de Acesso:", type="password")
    botao_entrar = st.button("Liberar Acesso 🔓")
    
    if botao_entrar:
        if senha_digitada == SENHA_CORRETA:
            st.session_state['autenticado'] = True
            st.rerun()
        else:
            st.error("⚠️ Senha incorreta. Verifique os dados ou adquira uma licença abaixo.")
            
    st.markdown("---")
    st.write("### 🚀 Ainda não tem uma senha de acesso?")
    st.write("Adquira sua chave de acesso imediata pelo Mercado Pago por apenas R$ 39,90.")
    
    # Botão visual chamativo para o Mercado Pago
    st.markdown(f'''
        <a href="{LINK_MERCADO_PAGO}" target="_blank">
            <button style="
                background-color: #009EE3; 
                color: white; 
                border: none; 
                padding: 12px 24px; 
                font-size: 18px; 
                font-weight: bold;
                border-radius: 5px; 
                cursor: pointer;
                width: 100%;">
                💳 Comprar Senha de Acesso no Mercado Pago
            </button>
        </a>
    ''', unsafe_allow_html=True)

# Tela do Sistema (Só carrega se o usuário estiver autenticado)
else:
    st.sidebar.success("🔑 Acesso Liberado")
    if st.sidebar.button("Sair / Bloquear"):
        st.session_state['autenticado'] = False
        st.rerun()

    st.write("Insira os dados do seu cliente abaixo para gerar a proposta comercial perfeita.")

    with st.form("dados_proposta"):
        nome_cliente = st.text_input("Nome do Cliente / Empresa:", placeholder="Ex: Clínica OdontoSorriso")
        servico = st.text_input("Serviço Ofertado:", placeholder="Ex: Gestão de Tráfego Pago e Redes Sociais")
        dores_cliente = st.text_area("Quais os problemas/dores que o cliente enfrenta hoje?", 
                                     placeholder="Ex: Dependem apenas de indicações e o faturamento travou.")
        preco_estimado = st.text_input("Valor do Investimento / Condições:", placeholder="Ex: R$ 2.500/mês")
        
        botao_gerar = st.form_submit_button("Gerar Proposta Comercial 🔥")

    def criar_arquivo_word(texto_proposta):
        doc = Document()
        doc.add_heading('Proposta Comercial', level=1)
        for linha in texto_proposta.split('\n'):
            if linha.strip().startswith('**') and linha.strip().endswith('**'):
                doc.add_heading(linha.replace('**', ''), level=2)
            elif linha.strip().startswith('#'):
                doc.add_heading(linha.replace('#', '').strip(), level=2)
            else:
                doc.add_paragraph(linha.replace('**', '').replace('*', ''))
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    if botao_gerar:
        if not nome_cliente or not servico or not dores_cliente or not preco_estimado:
            st.error("⚠️ Por favor, preencha todos os campos do formulário.")
        else:
            with st.spinner("Nossa Inteligência Artificial está estruturando sua proposta..."):
                try:
                    # Configura a API internamente com a sua chave oculta
                    genai.configure(api_key=MINHA_API_KEY_PRIVADA)
                    model = genai.GenerativeModel('gemini-3.5-flash')
                    
                    prompt = f"""
                    Você é um copywriter profissional e especialista em vendas B2B de alto impacto.
                    Escreva uma proposta comercial altamente persuasiva com base nos seguintes dados:
                    - Cliente: {nome_cliente}
                    - Serviço Prestado: {servico}
                    - Principais problemas/dores que o cliente enfrenta: {dores_cliente}
                    - Valor do Investimento: {preco_estimado}
                    
                    A proposta deve seguir a seguinte estrutura estrita em Markdown:
                    1. **Sumário Executivo**: Um gancho forte mostrando que entendemos a dor atual dele.
                    2. **A Solução**: Como o nosso serviço resolve especificamente as dores mencionadas.
                    3. **Escopo do Trabalho**: Entregáveis claros em tópicos.
                    4. **Investimento e Condições**: O valor de {preco_estimado} detalhado e profissionalmente justificado como investimento (ROI).
                    5. **Próximos Passos**: Chamada para ação clara para fechar o contrato.
                    
                    Use um tom profissional, corporativo, porém altamente focado em conversão e resultados. Não inclua saudações fora do Sumário.
                    """
                    response = model.generate_content(prompt)
                    st.session_state['proposta'] = response.text
                    st.session_state['nome_cliente'] = nome_cliente
                except Exception as e:
                    st.error(f"⚠️ Erro técnico real: {e}")


    if 'proposta' in st.session_state:
        st.success("✨ Proposta gerada com sucesso!")
        st.markdown("---")
        st.markdown(st.session_state['proposta'])
        st.markdown("---")
        
        arquivo_docx = criar_arquivo_word(st.session_state['proposta'])
        nome_arquivo = f"Proposta_{st.session_state['nome_cliente'].replace(' ', '_')}.docx"
        
        st.download_button(
            label="📥 Baixar Proposta em Word (.docx)",
            data=arquivo_docx,
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
